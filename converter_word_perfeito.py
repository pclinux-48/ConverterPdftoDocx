import sys
from pathlib import Path
import argparse

try:
    import pdfplumber
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re
except ImportError as e:
    print(f"Erro: Biblioteca necessária não encontrada - {e}")
    print("Execute: pip install pdfplumber python-docx")
    sys.exit(1)

class PDFToWordPerfeito:
    def __init__(self, pdf_path):
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"Arquivo PDF não encontrado: {pdf_path}")
    
    def _preserve_spacing(self, text):
        """Preserva espaçamento e formatação do texto"""
        if not text:
            return ""
        
        # Preservar múltiplos espaços
        text = re.sub(r' {2,}', lambda m: '\t' * (len(m.group()) // 2), text)
        return text
    
    def _add_formatted_paragraph(self, doc, text, font_size=11, bold=False, center=False):
        """Adiciona parágrafo com formatação específica"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        
        # Configurar fonte
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.bold = bold
        
        # Alinhamento
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return para
    
    def _create_table_from_data(self, doc, table_data, title=None):
        """Cria tabela no Word preservando formatação"""
        if not table_data:
            return
        
        if title:
            self._add_formatted_paragraph(doc, title, font_size=12, bold=True)
        
        # Filtrar linhas vazias
        filtered_data = []
        for row in table_data:
            if any(cell and str(cell).strip() for cell in row):
                filtered_data.append(row)
        
        if not filtered_data:
            return
        
        # Criar tabela
        table = doc.add_table(rows=len(filtered_data), cols=len(filtered_data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Preencher dados
        for row_idx, row_data in enumerate(filtered_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell_text = str(cell_data).strip() if cell_data else ""
                cell.text = cell_text
                
                # Formatação da célula
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                
                # Primeira linha em negrito (cabeçalho)
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajustar largura das colunas
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.5)
    
    def _extract_structured_content(self):
        """Extrai conteúdo de forma estruturada"""
        pages_content = []
        
        with pdfplumber.open(self.pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                print(f"Processando página {page_num}...")
                
                page_content = {
                    'page_num': page_num,
                    'text_blocks': [],
                    'tables': [],
                    'layout_preserved': False
                }
                
                # Estratégia 1: Extrair tabelas com linhas rígidas
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "lines_strict",
                    "horizontal_strategy": "lines_strict",
                    "snap_tolerance": 5,
                    "join_tolerance": 5,
                    "edge_min_length": 10,
                    "min_words_vertical": 1,
                    "min_words_horizontal": 1,
                })
                
                # Se não encontrar com linhas rígidas, tentar com linhas normais
                if not tables:
                    tables = page.extract_tables(table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                    })
                
                # Se não encontrar com linhas, tentar com texto
                if not tables:
                    tables = page.extract_tables(table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                    })
                
                # Última tentativa: configuração padrão
                if not tables:
                    tables = page.extract_tables()
                
                page_content['tables'] = tables
                
                # Estratégia 2: Extrair texto preservando layout
                text_layout = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                if text_layout:
                    # Dividir em blocos lógicos
                    lines = text_layout.split('\n')
                    current_block = []
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            current_block.append(line)
                        else:
                            if current_block:
                                page_content['text_blocks'].append('\n'.join(current_block))
                                current_block = []
                    
                    # Adicionar último bloco
                    if current_block:
                        page_content['text_blocks'].append('\n'.join(current_block))
                    
                    page_content['layout_preserved'] = True
                
                # Estratégia 3: Se não conseguiu preservar layout, extrair texto normal
                if not page_content['layout_preserved']:
                    text_normal = page.extract_text()
                    if text_normal:
                        page_content['text_blocks'] = [text_normal]
                
                pages_content.append(page_content)
        
        return pages_content
    
    def convert_to_word(self, output_path=None):
        """Converte PDF para Word com máxima fidelidade"""
        if output_path is None:
            base_name = self.pdf_path.stem
            output_path = self.pdf_path.parent / f"{base_name}_perfeito.docx"
        
        print(f"🚀 Iniciando conversão perfeita para Word...")
        print(f"📄 Arquivo origem: {self.pdf_path}")
        print(f"📄 Arquivo destino: {output_path}")
        
        # Extrair conteúdo estruturado
        pages_content = self._extract_structured_content()
        
        # Criar documento Word
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar título
        title = self._add_formatted_paragraph(
            doc, 
            f"Conversão de: {self.pdf_path.name}", 
            font_size=14, 
            bold=True, 
            center=True
        )
        doc.add_paragraph()  # Espaço
        
        # Processar cada página
        for page_content in pages_content:
            page_num = page_content['page_num']
            
            # Adicionar cabeçalho da página (se mais de uma página)
            if len(pages_content) > 1:
                self._add_formatted_paragraph(
                    doc, 
                    f"PÁGINA {page_num}", 
                    font_size=12, 
                    bold=True, 
                    center=True
                )
                doc.add_paragraph()
            
            # Adicionar tabelas primeiro
            if page_content['tables']:
                print(f"  📊 Processando {len(page_content['tables'])} tabela(s) da página {page_num}")
                
                for table_idx, table_data in enumerate(page_content['tables']):
                    if table_data:
                        table_title = f"Tabela {table_idx + 1}" if len(page_content['tables']) > 1 else None
                        self._create_table_from_data(doc, table_data, table_title)
                        doc.add_paragraph()  # Espaço após tabela
            
            # Adicionar blocos de texto
            if page_content['text_blocks']:
                print(f"  📝 Processando {len(page_content['text_blocks'])} bloco(s) de texto da página {page_num}")
                
                for block in page_content['text_blocks']:
                    if block.strip():
                        # Verificar se é um título (linha curta, maiúsculas, etc.)
                        is_title = (
                            len(block.strip()) < 100 and 
                            (block.isupper() or block.count(' ') < 5)
                        )
                        
                        if is_title:
                            self._add_formatted_paragraph(doc, block, font_size=12, bold=True, center=True)
                        else:
                            formatted_text = self._preserve_spacing(block)
                            self._add_formatted_paragraph(doc, formatted_text, font_size=11)
                        
                        doc.add_paragraph()  # Espaço entre blocos
            
            # Quebra de página (exceto na última página)
            if page_num < len(pages_content):
                doc.add_page_break()
        
        # Salvar documento
        doc.save(output_path)
        
        print(f"✅ Conversão concluída com sucesso!")
        print(f"📄 Arquivo Word salvo em: {output_path}")
        
        return output_path

def main():
    parser = argparse.ArgumentParser(description='Conversor PDF para Word Perfeito')
    parser.add_argument('pdf_file', help='Caminho para o arquivo PDF')
    parser.add_argument('-o', '--output', help='Caminho do arquivo Word de saída')
    
    args = parser.parse_args()
    
    try:
        converter = PDFToWordPerfeito(args.pdf_file)
        output_file = converter.convert_to_word(args.output)
        
        print(f"\n🎉 Conversão perfeita concluída!")
        print(f"📂 Abra o arquivo: {output_file}")
        
    except Exception as e:
        print(f"❌ Erro durante a conversão: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) == 1:
        # Procurar PDF na pasta atual
        pdf_files = list(Path('.').glob('*.pdf'))
        if pdf_files:
            print(f"📄 Convertendo: {pdf_files[0]}")
            converter = PDFToWordPerfeito(pdf_files[0])
            converter.convert_to_word()
        else:
            print("❌ Nenhum arquivo PDF encontrado.")
            print("💡 Use: python converter_word_perfeito.py 'arquivo.pdf'")
    else:
        main()