import sys
from pathlib import Path
import argparse

# Importações e Classes de Formatação (mantidas as originais)
try:
    import pdfplumber
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re
except ImportError as e:
    print(f"Erro: Biblioteca necessária não encontrada - {e}")
    print("Execute: pip install pdfplumber python-docx")
    sys.exit(1)

# Defina o tamanho do lote de conversão
TAMANHO_DO_LOTE = 50

class PDFToWordPerfeito:
    def __init__(self, pdf_path):
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"Arquivo PDF não encontrado: {pdf_path}")
        
    # --- Métodos Auxiliares de Formatação (Mantidos) ---
    def _preserve_spacing(self, text):
        if not text: return ""
        text = re.sub(r' {2,}', lambda m: ' ' * (len(m.group()) // 2 + 1), text)
        return text
    
    def _add_formatted_paragraph(self, doc, text, font_size=11, bold=False, center=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.bold = bold
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para
    
    def _create_table_from_data(self, doc, table_data, title=None):
        if not table_data: return
        if title: self._add_formatted_paragraph(doc, title, font_size=12, bold=True)
        filtered_data = [row for row in table_data if any(cell and str(cell).strip() for cell in row)]
        if not filtered_data: return
        
        table = doc.add_table(rows=len(filtered_data), cols=len(filtered_data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row_data in enumerate(filtered_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell_text = str(cell_data).strip() if cell_data else ""
                cell.text = cell_text
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.5)

    def _extract_page_content(self, page, page_num):
        """Extrai conteúdo estruturado de uma ÚNICA página."""
        page_content = {
            'page_num': page_num,
            'text_blocks': [],
            'tables': [],
            'layout_preserved': False
        }
        
        # --- LÓGICA DE EXTRAÇÃO DE TABELAS ---
        # Tenta várias estratégias para extrair tabelas
        table_settings = [
            {"vertical_strategy": "lines_strict", "horizontal_strategy": "lines_strict", "snap_tolerance": 5, "join_tolerance": 5, "edge_min_length": 10, "min_words_vertical": 1, "min_words_horizontal": 1},
            {"vertical_strategy": "lines", "horizontal_strategy": "lines", "snap_tolerance": 3, "join_tolerance": 3},
            {"vertical_strategy": "text", "horizontal_strategy": "text", "snap_tolerance": 3, "join_tolerance": 3},
            {} # Padrão
        ]
        
        tables = []
        for settings in table_settings:
            tables = page.extract_tables(table_settings=settings)
            if tables and any(tables):
                break
        
        page_content['tables'] = tables
        
        # --- LÓGICA DE EXTRAÇÃO DE TEXTO ---
        text_layout = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
        if text_layout:
            lines = text_layout.split('\n')
            current_block = []
            
            for line in lines:
                line = line.strip()
                if line:
                    current_block.append(line)
                else:
                    if current_block:
                        page_content['text_blocks'].append(' '.join(current_block))
                        current_block = []
            
            if current_block:
                page_content['text_blocks'].append(' '.join(current_block))
            
            page_content['layout_preserved'] = True
        
        if not page_content['layout_preserved']:
            text_normal = page.extract_text()
            if text_normal:
                page_content['text_blocks'] = [re.sub(r'\n{2,}', '\n', text_normal)]
                
        return page_content

    def _process_page_content(self, doc, page_content, is_first_page_in_batch=False):
        """Processa o conteúdo de uma única página e adiciona ao docx"""
        page_num = page_content['page_num']
        
        # Adicionar quebra de página se NÃO for a primeira do lote
        if not is_first_page_in_batch:
            doc.add_page_break()

        # Adicionar cabeçalho da página
        self._add_formatted_paragraph(
            doc, 
            f"PÁGINA {page_num}", 
            font_size=12, 
            bold=True, 
            center=True
        )
        doc.add_paragraph()
        
        # Adicionar tabelas primeiro
        if page_content['tables']:
            for table_idx, table_data in enumerate(page_content['tables']):
                if table_data:
                    table_title = f"Tabela {table_idx + 1}" if len(page_content['tables']) > 1 else None
                    self._create_table_from_data(doc, table_data, table_title)
                    doc.add_paragraph()
        
        # Adicionar blocos de texto
        if page_content['text_blocks']:
            for block in page_content['text_blocks']:
                if block.strip():
                    is_title = (
                        len(block.strip()) < 100 and 
                        (block.isupper() or block.count(' ') < 5) and
                        not block.startswith(' ')
                    )
                    
                    if is_title:
                        self._add_formatted_paragraph(doc, block, font_size=12, bold=True, center=True)
                    else:
                        formatted_text = self._preserve_spacing(block)
                        self._add_formatted_paragraph(doc, formatted_text, font_size=11)
                    
                    doc.add_paragraph()

    def convert_to_word(self, output_path=None):
        """Converte PDF para Word em lotes de 50 páginas, extraindo página a página."""
        
        print(f"🚀 Iniciando conversão perfeita para Word em lotes de {TAMANHO_DO_LOTE} páginas...")
        print(f"📄 Arquivo origem: {self.pdf_path}")
        
        base_name = self.pdf_path.stem
        total_pages = 0
        lote_num = 0

        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                total_pages = len(pdf.pages)
                
                # Loop para processar e salvar em lotes
                for start_index in range(0, total_pages, TAMANHO_DO_LOTE):
                    end_index = min(start_index + TAMANHO_DO_LOTE, total_pages)
                    lote_num += 1
                    
                    # Define o nome do arquivo de saída para o lote
                    output_file_lote = self.pdf_path.parent / f"{base_name}_parte_{lote_num:02d}.docx"
                    
                    doc = Document()
                    print(f"\n📂 Processando LOTE {lote_num}: Páginas {start_index + 1} a {end_index}...")

                    # Configurar margens (uma vez por documento)
                    for section in doc.sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                        
                    # Adicionar título no início de cada documento
                    self._add_formatted_paragraph(
                        doc, 
                        f"Conversão de: {self.pdf_path.name} (Parte {lote_num})", 
                        font_size=14, 
                        bold=True, 
                        center=True
                    )
                    doc.add_paragraph()

                    # Processar e extrair DENTRO do loop do lote
                    for page_index in range(start_index, end_index):
                        page_num_real = page_index + 1
                        page = pdf.pages[page_index]
                        
                        print(f"  -> Extraindo e processando página {page_num_real} de {total_pages}...")
                        
                        # Extrai o conteúdo da página atual
                        page_content = self._extract_page_content(page, page_num_real)
                        
                        is_first = (page_index == start_index)
                        
                        # Adiciona o conteúdo extraído ao documento Word
                        self._process_page_content(doc, page_content, is_first_page_in_batch=is_first)
                    
                    # Salvar documento do lote
                    doc.save(output_file_lote)
                    print(f"✅ Lote {lote_num} concluído e salvo em: {output_file_lote.name}")
            
            print(f"\n🎉 Conversão em lotes concluída! Total de {total_pages} páginas processadas.")
            return True

        except Exception as e:
            print(f"❌ Erro fatal durante a conversão do lote {lote_num}: {e}")
            import traceback
            traceback.print_exc()
            return False


def main():
    parser = argparse.ArgumentParser(description='Conversor PDF para Word Perfeito')
    parser.add_argument('pdf_file', nargs='?', default=None, help='Caminho/Nome do arquivo PDF')
    parser.add_argument('-o', '--output', help='Caminho do arquivo Word de saída (Ignorado no modo Lote)')
    
    args = parser.parse_args()
    
    try:
        if args.pdf_file:
            pdf_path = Path(args.pdf_file)
            if not pdf_path.is_absolute():
                pdf_path = Path.cwd() / args.pdf_file
            
            converter = PDFToWordPerfeito(pdf_path)
            converter.convert_to_word(args.output)
            
        else:
            pdf_files = list(Path('.').glob('*.pdf'))
            if pdf_files:
                print(f"📄 Convertendo o primeiro PDF encontrado: {pdf_files[0].name}")
                converter = PDFToWordPerfeito(pdf_files[0])
                converter.convert_to_word()
            else:
                print("❌ Nenhum arquivo PDF encontrado na pasta atual.")
                print("💡 Use: python seu_script.py 'nome_do_arquivo.pdf'")

    except Exception as e:
        print(f"❌ Erro durante a inicialização: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()