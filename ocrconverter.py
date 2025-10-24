import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from pdf2image import convert_from_path 
import pytesseract                     

# --- CONFIGURAÇÕES CRÍTICAS DO TESSERACT PARA SEU AMBIENTE macOS ---
# Estes caminhos foram verificados no seu sistema (Homebrew).
TESSERACT_EXECUTABLE_MAC = '/opt/homebrew/bin/tesseract'
TESSDATA_PREFIX_MAC = '/opt/homebrew/share/tessdata' # O diretório que contém por.traineddata
IDIOMA_OCR = 'por' 

# Define o tamanho do lote de páginas para cada arquivo DOCX
TAMANHO_DO_LOTE = 50 

def converter_pdf_com_ocr_em_lotes(nome_arquivo_pdf):
    """
    Converte um PDF baseado em imagem para DOCX usando OCR em lotes,
    passando os caminhos de configuração diretamente ao pytesseract.
    """
    
    # Define caminhos
    diretorio_atual = os.path.dirname(os.path.abspath(__file__))
    caminho_pdf = os.path.join(diretorio_atual, nome_arquivo_pdf)

    if not os.path.exists(caminho_pdf):
        print(f"\nERRO: Arquivo PDF não encontrado em: {caminho_pdf}")
        return

    nome_base = os.path.splitext(os.path.basename(caminho_pdf))[0]
    caminho_saida_base = os.path.join(diretorio_atual, nome_base)

    print(f"\n🚀 Iniciando conversão via OCR (Tesseract) em lotes de {TAMANHO_DO_LOTE} páginas...")
    print("ATENÇÃO: Este processo é mais lento, mas necessário para PDFs baseados em imagem.")
    
    try:
        # **CORREÇÃO DEFINITIVA:**
        # 1. Define o caminho do executável do Tesseract.
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXECUTABLE_MAC
        
        # 2. Configuração de idioma passada como parâmetro customizado,
        #    forçando o Tesseract a procurar os dados no local correto.
        tesseract_config = f'--tessdata-dir "{TESSDATA_PREFIX_MAC}"'
        
        # 1. Converte o PDF inteiro em uma lista de objetos imagem
        print("\nConvertendo PDF para imagens (Requer Poppler instalado)...")
        pages_images = convert_from_path(caminho_pdf)
        total_paginas = len(pages_images)
        lote_atual = 1
        
        # 2. Processamento em lotes
        for inicio_pagina in range(0, total_paginas, TAMANHO_DO_LOTE):
            
            fim_pagina = min(inicio_pagina + TAMANHO_DO_LOTE, total_paginas)
            caminho_docx_lote = f"{caminho_saida_base}_OCR_parte_{lote_atual:02d}.docx"
            documento_word = Document()
            
            print(f"\n📂 Processando LOTE {lote_atual}: Páginas {inicio_pagina+1} a {fim_pagina} via OCR...")

            # Configurar margens
            for section in documento_word.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # 3. Itera pelas imagens do lote e aplica OCR
            for num_pagina in range(inicio_pagina, fim_pagina):
                
                # Adiciona quebra de página se não for a primeira do lote
                if num_pagina > inicio_pagina:
                    documento_word.add_page_break()
                    
                imagem_pagina = pages_images[num_pagina]
                
                # Realiza o OCR, passando a configuração customizada (caminho dos dados)
                texto_pagina = pytesseract.image_to_string(
                    imagem_pagina, 
                    lang=IDIOMA_OCR,
                    config=tesseract_config
                )
                
                # Adiciona cabeçalho da página
                paragrafo_cabecalho = documento_word.add_paragraph()
                paragrafo_cabecalho.add_run(f"--- PÁGINA {num_pagina+1} ---").bold = True
                paragrafo_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adiciona o texto extraído
                if texto_pagina and texto_pagina.strip():
                    documento_word.add_paragraph(texto_pagina)
                    print(f"   Página {num_pagina+1}: Texto extraído com sucesso.")
                else:
                    documento_word.add_paragraph("[AVISO: Nenhum texto reconhecido nesta página, ou página em branco.]")
                    print(f"   Página {num_pagina+1}: Falha na extração de texto (pode ser imagem sem texto ou ilegível).")
                
            # 4. Salva o documento DOCX do lote
            documento_word.save(caminho_docx_lote)
            print(f"✅ LOTE {lote_atual} CONCLUÍDO e salvo em: {os.path.basename(caminho_docx_lote)}")
            
            lote_atual += 1
            
        print("\n====================================================================")
        print(f"CONVERSÃO OCR COMPLETA! Total de {total_paginas} páginas convertidas.")
        print("====================================================================")
    
    except pytesseract.TesseractNotFoundError:
        print("\nERRO FATAL: Tesseract OCR não encontrado. (Verifique o caminho: /opt/homebrew/bin/tesseract)")
    except Exception as e:
        print(f"\n❌ ERRO grave durante a conversão: {e}")
        print("Verifique se o Poppler está instalado corretamente (necessário para pdf2image) e no PATH.")
        import traceback
        traceback.print_exc()


# --- EXECUÇÃO PRINCIPAL ---
if __name__ == "__main__":
    
    # Obtém o nome do arquivo PDF através da entrada do usuário
    nome_arquivo_pdf = input("Por favor, digite o NOME COMPLETO do arquivo PDF (ex: relatorio.pdf): ")
    
    # Chama a função de conversão
    converter_pdf_com_ocr_em_lotes(nome_arquivo_pdf)